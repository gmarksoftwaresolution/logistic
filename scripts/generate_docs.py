import sys
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# --- PATHS ---
DOCS_DIR = r"c:\Users\parth\OneDrive\Desktop\GST-v1\logistic\docs"
MD_FILE = os.path.join(DOCS_DIR, "GMU_Logistics_Master_Onboarding_Guide.md")
HTML_FILE = os.path.join(DOCS_DIR, "GMU_Logistics_Master_Onboarding_Guide.html")
DOCX_FILE = os.path.join(DOCS_DIR, "GMU_Logistics_Master_Onboarding_Guide.docx")

def read_markdown():
    with open(MD_FILE, "r", encoding="utf-8") as f:
        return f.read()

# --- HTML GENERATOR ---
def generate_html(md_content):
    # Parse markdown blocks into HTML structure
    lines = md_content.splitlines()
    html_out = []
    
    in_code_block = False
    code_lang = ""
    code_buffer = []
    
    in_table = False
    table_buffer = []
    
    in_callout = False
    callout_type = "info"
    callout_buffer = []
    
    in_list = False
    
    toc_items = []
    
    # First pass: collect TOC headings
    for line in lines:
        if line.startswith("# "):
            title = line[2:].strip()
            toc_items.append((1, title, re.sub(r'[^a-z0-9\-]', '', title.lower().replace(' ', '-'))))
        elif line.startswith("## "):
            title = line[3:].strip()
            if not title.startswith("Complete Project Documentation"):
                toc_items.append((2, title, re.sub(r'[^a-z0-9\-]', '', title.lower().replace(' ', '-'))))

    html_out.append('''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>GMU Logistics Platform - Master Onboarding Guide</title>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Fira+Code:wght@400;500;600&family=Poppins:ital,wght@0,300;0,400;0,500;0,600;0,700;1,400&display=swap" rel="stylesheet">
    <style>
        /* === PRINT & PAGE LAYOUT SETUP === */
        @page {
            size: A4;
            margin: 20mm 15mm 20mm 15mm;
            @bottom-right {
                content: "Page " counter(page) " of " counter(pages);
                font-family: 'Poppins', sans-serif;
                font-size: 8pt;
                color: #64748b;
            }
            @bottom-left {
                content: "Gramuunati Logistics (GMU) Platform";
                font-family: 'Poppins', sans-serif;
                font-size: 8pt;
                color: #64748b;
            }
        }

        * {
            box-sizing: border-box;
            -webkit-print-color-adjust: exact !important;
            print-color-adjust: exact !important;
        }

        body {
            font-family: 'Poppins', sans-serif;
            color: #0f172a;
            background-color: #ffffff;
            line-height: 1.6;
            font-size: 10pt;
            margin: 0;
            padding: 0;
        }

        /* === COVER PAGE === */
        .cover-page {
            height: 90vh;
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            text-align: center;
            page-break-after: always;
            break-after: page;
            background: linear-gradient(135deg, #f0fdf4 0%, #ffffff 100%);
            border: 2px solid #bbf7d0;
            border-radius: 12px;
            padding: 40px;
            margin-bottom: 40px;
        }

        .cover-badge {
            background-color: #15803d;
            color: #ffffff;
            font-size: 9pt;
            font-weight: 600;
            padding: 6px 16px;
            border-radius: 20px;
            text-transform: uppercase;
            letter-spacing: 1px;
            margin-bottom: 24px;
        }

        .cover-title {
            font-size: 28pt;
            font-weight: 700;
            color: #15803d;
            margin: 0 0 12px 0;
            line-height: 1.2;
        }

        .cover-subtitle {
            font-size: 16pt;
            font-weight: 500;
            color: #334155;
            margin: 0 0 40px 0;
        }

        .cover-meta {
            margin-top: 40px;
            border-top: 2px solid #e2e8f0;
            padding-top: 24px;
            width: 80%;
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 16px;
            text-align: left;
            font-size: 9.5pt;
        }

        .meta-item strong {
            color: #166534;
        }

        /* === HEADINGS & SECTIONS === */
        h1.section-title {
            font-size: 18pt;
            font-weight: 700;
            color: #15803d;
            border-bottom: 2px solid #bbf7d0;
            padding-bottom: 8px;
            margin-top: 36px;
            margin-bottom: 16px;
            page-break-before: always;
            break-before: page;
        }

        h2 {
            font-size: 14pt;
            font-weight: 600;
            color: #166534;
            margin-top: 24px;
            margin-bottom: 12px;
            break-after: avoid;
        }

        h3 {
            font-size: 11.5pt;
            font-weight: 600;
            color: #0f172a;
            margin-top: 18px;
            margin-bottom: 8px;
            break-after: avoid;
        }

        p {
            margin-top: 0;
            margin-bottom: 12px;
            text-align: justify;
        }

        /* === TABLE OF CONTENTS === */
        .toc-container {
            background-color: #f8fafc;
            border: 1px solid #e2e8f0;
            border-radius: 8px;
            padding: 24px;
            margin-bottom: 32px;
            page-break-after: always;
            break-after: page;
        }

        .toc-title {
            font-size: 16pt;
            font-weight: 700;
            color: #15803d;
            margin-top: 0;
            margin-bottom: 16px;
            border-bottom: 2px solid #15803d;
            padding-bottom: 6px;
        }

        .toc-list {
            list-style: none;
            padding-left: 0;
            margin: 0;
        }

        .toc-item-1 {
            font-weight: 600;
            font-size: 10pt;
            margin-top: 10px;
            color: #15803d;
        }

        .toc-item-2 {
            font-weight: 400;
            font-size: 9pt;
            margin-left: 20px;
            margin-top: 4px;
            color: #334155;
        }

        .toc-link {
            color: inherit;
            text-decoration: none;
        }

        .toc-link:hover {
            text-decoration: underline;
        }

        /* === CALLOUT BOXES === */
        .callout {
            border-left: 4px solid #15803d;
            background-color: #f0fdf4;
            padding: 14px 18px;
            border-radius: 0 8px 8px 0;
            margin: 16px 0;
            break-inside: avoid;
            font-size: 9.5pt;
        }

        .callout-important {
            border-left-color: #15803d;
            background-color: #f0fdf4;
        }

        .callout-warning {
            border-left-color: #d97706;
            background-color: #fffbeb;
        }

        .callout-tip {
            border-left-color: #0284c7;
            background-color: #f0f9ff;
        }

        .callout-note {
            border-left-color: #64748b;
            background-color: #f8fafc;
        }

        .callout-title {
            font-weight: 700;
            margin-bottom: 6px;
            display: flex;
            align-items: center;
            gap: 6px;
        }

        .callout-important .callout-title { color: #15803d; }
        .callout-warning .callout-title { color: #b45309; }
        .callout-tip .callout-title { color: #0369a1; }
        .callout-note .callout-title { color: #475569; }

        /* === TABLES === */
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 18px 0;
            font-size: 9pt;
            break-inside: avoid;
        }

        th {
            background-color: #15803d;
            color: #ffffff;
            font-weight: 600;
            text-align: left;
            padding: 10px 12px;
            border: 1px solid #15803d;
        }

        td {
            padding: 8px 12px;
            border: 1px solid #cbd5e1;
            vertical-align: top;
        }

        tr:nth-child(even) {
            background-color: #f8fafc;
        }

        /* === CODE BLOCKS & DIAGRAMS === */
        pre {
            background-color: #0f172a;
            color: #f8fafc;
            font-family: 'Fira Code', monospace;
            font-size: 8.5pt;
            padding: 14px 16px;
            border-radius: 6px;
            overflow-x: auto;
            margin: 16px 0;
            break-inside: avoid;
            line-height: 1.45;
            border: 1px solid #1e293b;
        }

        code {
            font-family: 'Fira Code', monospace;
            background-color: #f1f5f9;
            color: #0f172a;
            padding: 2px 6px;
            border-radius: 4px;
            font-size: 8.5pt;
            border: 1px solid #e2e8f0;
        }

        pre code {
            background-color: transparent;
            color: inherit;
            padding: 0;
            border: none;
        }

        .mermaid-box {
            background-color: #f8fafc;
            border: 1px solid #cbd5e1;
            border-left: 4px solid #15803d;
            border-radius: 6px;
            padding: 16px;
            margin: 16px 0;
            font-family: 'Fira Code', monospace;
            font-size: 8.5pt;
            white-space: pre-wrap;
            break-inside: avoid;
            color: #1e293b;
        }

        /* === LISTS === */
        ul, ol {
            margin-top: 0;
            margin-bottom: 12px;
            padding-left: 24px;
        }

        li {
            margin-bottom: 4px;
        }

        /* === FOOTER PRINT === */
        .page-footer {
            margin-top: 40px;
            padding-top: 12px;
            border-top: 1px solid #e2e8f0;
            font-size: 8pt;
            color: #64748b;
            display: flex;
            justify-content: space-between;
        }
    </style>
</head>
<body>
''')

    # Add Cover Page
    html_out.append('''
    <div class="cover-page">
        <div class="cover-badge">Gramuunati Logistics Monorepo</div>
        <h1 class="cover-title">GMU Logistics Platform</h1>
        <div class="cover-subtitle">Complete Project Documentation & Master Onboarding Guide</div>
        
        <div class="cover-meta">
            <div class="meta-item"><strong>Document Version:</strong> 1.0.0</div>
            <div class="meta-item"><strong>Company:</strong> Gramuunati Logistics (GMU)</div>
            <div class="meta-item"><strong>Target Audience:</strong> Interns, Developers, QA & New Team Members</div>
            <div class="meta-item"><strong>Published:</strong> August 2026</div>
        </div>
    </div>
    ''')

    # Add Table of Contents
    html_out.append('<div class="toc-container"><div class="toc-title">Table of Contents</div><ul class="toc-list">')
    for level, title, anchor in toc_items:
        cls = f"toc-item-{level}"
        html_out.append(f'<li class="{cls}"><a href="#{anchor}" class="toc-link">{title}</a></li>')
    html_out.append('</ul></div>')

    # Process Markdown lines
    def format_inline(text):
        # Format inline bold, italic, code, links
        text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'\*(.*?)\*', r'<em>\1</em>', text)
        text = re.sub(r'`(.*?)`', r'<code>\1</code>', text)
        text = re.sub(r'\[(.*?)\]\((.*?)\)', r'<a href="\2" target="_blank">\1</a>', text)
        return text

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block
        if line.startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_buffer)
                if code_lang in ["mermaid", "text"]:
                    html_out.append(f'<div class="mermaid-box"><strong>[Diagram / Workflow Architecture]</strong>\n{code_text}</div>')
                else:
                    html_out.append(f'<pre><code>{code_text}</code></pre>')
                in_code_block = False
                code_buffer = []
            else:
                in_code_block = True
                code_lang = line[3:].strip().lower()
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Callouts
        if line.startswith("> [!"):
            m = re.match(r'> \[!(IMPORTANT|WARNING|TIP|NOTE)\]', line)
            if m:
                in_callout = True
                callout_type = m.group(1).lower()
                callout_buffer = []
                i += 1
                continue

        if in_callout:
            if line.startswith("> "):
                callout_buffer.append(line[2:])
                i += 1
                continue
            else:
                c_text = format_inline(" ".join(callout_buffer))
                html_out.append(f'<div class="callout callout-{callout_type}"><div class="callout-title">📌 {callout_type.upper()}</div><div>{c_text}</div></div>')
                in_callout = False
                callout_buffer = []

        # Tables
        if "|" in line and not line.startswith("```"):
            if not in_table:
                in_table = True
                table_buffer = []
            table_buffer.append(line)
            i += 1
            continue
        elif in_table:
            # Process table
            if len(table_buffer) >= 2:
                html_out.append('<table>')
                headers = [c.strip() for c in table_buffer[0].split('|')[1:-1]]
                html_out.append('<thead><tr>' + ''.join([f'<th>{format_inline(h)}</th>' for h in headers]) + '</tr></thead><tbody>')
                
                for r_line in table_buffer[2:]:
                    cells = [c.strip() for c in r_line.split('|')[1:-1]]
                    if cells:
                        html_out.append('<tr>' + ''.join([f'<td>{format_inline(c)}</td>' for c in cells]) + '</tr>')
                html_out.append('</tbody></table>')
            in_table = False
            table_buffer = []

        # Headings
        if line.startswith("# "):
            htitle = line[2:].strip()
            anchor = re.sub(r'[^a-z0-9\-]', '', htitle.lower().replace(' ', '-'))
            html_out.append(f'<h1 class="section-title" id="{anchor}">{htitle}</h1>')
            i += 1
            continue
        elif line.startswith("## "):
            htitle = line[3:].strip()
            anchor = re.sub(r'[^a-z0-9\-]', '', htitle.lower().replace(' ', '-'))
            html_out.append(f'<h2 id="{anchor}">{htitle}</h2>')
            i += 1
            continue
        elif line.startswith("### "):
            htitle = line[4:].strip()
            html_out.append(f'<h3>{htitle}</h3>')
            i += 1
            continue

        # Lists
        if line.startswith("* ") or line.startswith("- ") or re.match(r'^\d+\.\s', line):
            if not in_list:
                in_list = True
                html_out.append('<ul>')
            item_text = re.sub(r'^(\*|-|\d+\.)\s+', '', line)
            html_out.append(f'<li>{format_inline(item_text)}</li>')
            i += 1
            continue
        elif in_list:
            html_out.append('</ul>')
            in_list = False

        # Paragraphs
        if line.strip():
            html_out.append(f'<p>{format_inline(line)}</p>')

        i += 1

    if in_list:
        html_out.append('</ul>')

    html_out.append('</body></html>')

    with open(HTML_FILE, "w", encoding="utf-8") as f:
        f.write("\n".join(html_out))
    print(f"Generated HTML: {HTML_FILE}")

# --- DOCX GENERATOR ---
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generate_docx(md_content):
    doc = Document()

    # Page setup A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(15, 23, 42)

    # Header and Footer
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Gramuunati Logistics (GMU) Monorepo Platform - Master Onboarding Guide"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.style.font.size = Pt(8.5)
    hp.style.font.color.rgb = RGBColor(100, 116, 139)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "GMU Logistics Engineering Team Documentation | Version 1.0.0"
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.style.font.size = Pt(8.5)
    fp.style.font.color.rgb = RGBColor(100, 116, 139)

    # --- COVER PAGE ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(80)
    title_run = title_p.add_run("GMU Logistics Platform")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(21, 128, 61) # Green 700

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(40)
    sub_run = sub_p.add_run("Complete Project Documentation & Master Onboarding Guide")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(71, 85, 105)

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Document Version:", "1.0.0"),
        ("Company:", "Gramuunati Logistics (GMU)"),
        ("Target Audience:", "Interns, Developers, QA & New Team Members"),
        ("Published:", "August 2026")
    ]
    for row_idx, (label, val) in enumerate(meta_data):
        r = meta_table.rows[row_idx]
        r.cells[0].paragraphs[0].add_run(label).bold = True
        r.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(22, 101, 52)
        r.cells[1].paragraphs[0].add_run(val)
        set_cell_background(r.cells[0], "F0FDF4")
        set_cell_background(r.cells[1], "F0FDF4")
        set_cell_margins(r.cells[0], 120, 120, 150, 150)
        set_cell_margins(r.cells[1], 120, 120, 150, 150)

    doc.add_page_break()

    # --- PARSE MARKDOWN INTO DOCX ---
    lines = md_content.splitlines()
    in_code = False
    code_buf = []
    in_table = False
    table_buf = []
    in_callout = False
    callout_type = "info"
    callout_buf = []

    def add_formatted_p(text, style_name='Normal', space_before=0, space_after=6, bold_prefix=None):
        p = doc.add_paragraph(style=style_name)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15

        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.bold = True

        # Simple inline formatting parsing
        parts = re.split(r'(\*\*.*?\*\*|`.*?`|\*.*?\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2])
                r.bold = True
            elif part.startswith('*') and part.endswith('*'):
                r = p.add_run(part[1:-1])
                r.italic = True
            elif part.startswith('`') and part.endswith('`'):
                r = p.add_run(part[1:-1])
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(15, 23, 42)
            else:
                p.add_run(part)
        return p

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block
        if line.startswith("```"):
            if in_code:
                c_text = "\n".join(code_buf)
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = tbl.cell(0, 0)
                set_cell_background(cell, "0F172A")
                set_cell_margins(cell, 150, 150, 200, 200)
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_before = Pt(4)
                cp.paragraph_format.space_after = Pt(4)
                c_run = cp.add_run(c_text)
                c_run.font.name = 'Consolas'
                c_run.font.size = Pt(8.5)
                c_run.font.color.rgb = RGBColor(248, 250, 252)
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
                in_code = False
                code_buf = []
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Callout
        if line.startswith("> [!"):
            m = re.match(r'> \[!(IMPORTANT|WARNING|TIP|NOTE)\]', line)
            if m:
                in_callout = True
                callout_type = m.group(1).lower()
                callout_buf = []
                i += 1
                continue

        if in_callout:
            if line.startswith("> "):
                callout_buf.append(line[2:])
                i += 1
                continue
            else:
                c_text = " ".join(callout_buf)
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = tbl.cell(0, 0)
                bg_color = "F0FDF4" if callout_type == "important" else ("FFFBEB" if callout_type == "warning" else "F8FAFC")
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, 140, 140, 180, 180)
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_after = Pt(2)
                t_run = cp.add_run(f"📌 {callout_type.upper()}: ")
                t_run.bold = True
                t_run.font.color.rgb = RGBColor(21, 128, 61) if callout_type == "important" else RGBColor(180, 83, 9)
                add_formatted_p(c_text, space_after=2)
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
                in_callout = False
                callout_buf = []

        # Table
        if "|" in line and not line.startswith("```"):
            if not in_table:
                in_table = True
                table_buf = []
            table_buf.append(line)
            i += 1
            continue
        elif in_table:
            if len(table_buf) >= 2:
                headers = [c.strip() for c in table_buf[0].split('|')[1:-1]]
                rows_data = []
                for r_line in table_buf[2:]:
                    cells = [c.strip() for c in r_line.split('|')[1:-1]]
                    if cells:
                        rows_data.append(cells)

                tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Header Row
                hdr_cells = tbl.rows[0].cells
                for idx, h_text in enumerate(headers):
                    hdr_cells[idx].text = h_text
                    set_cell_background(hdr_cells[idx], "15803D")
                    set_cell_margins(hdr_cells[idx], 120, 120, 140, 140)
                    p = hdr_cells[idx].paragraphs[0]
                    p.runs[0].font.bold = True
                    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    p.runs[0].font.size = Pt(9)

                # Data Rows
                for r_idx, r_data in enumerate(rows_data):
                    row_cells = tbl.rows[r_idx + 1].cells
                    bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
                    for c_idx, c_val in enumerate(r_data):
                        if c_idx < len(row_cells):
                            row_cells[c_idx].text = c_val
                            set_cell_background(row_cells[c_idx], bg)
                            set_cell_margins(row_cells[c_idx], 100, 100, 120, 120)
                            p = row_cells[c_idx].paragraphs[0]
                            p.style.font.size = Pt(8.5)

                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            in_table = False
            table_buf = []

        # Headings
        if line.startswith("# "):
            doc.add_page_break()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[2:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(21, 128, 61)
            i += 1
            continue
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[3:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(22, 101, 52)
            i += 1
            continue
        elif line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[4:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(15, 23, 42)
            i += 1
            continue

        # Bullet List
        if line.startswith("* ") or line.startswith("- ") or re.match(r'^\d+\.\s', line):
            item_text = re.sub(r'^(\*|-|\d+\.)\s+', '', line)
            p = add_formatted_p(item_text, space_before=0, space_after=3)
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue

        # Paragraph
        if line.strip():
            add_formatted_p(line)

        i += 1

    try:
        doc.save(DOCX_FILE)
        print(f"Generated DOCX: {DOCX_FILE}")
    except PermissionError:
        alt_docx = DOCX_FILE.replace(".docx", "_v1.docx")
        doc.save(alt_docx)
        print(f"Generated DOCX (Primary file locked by MS Word, saved to): {alt_docx}")


if __name__ == "__main__":
    content = read_markdown()
    generate_html(content)
    generate_docx(content)
